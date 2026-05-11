"""
Generates the VU CS619 Final Project Report as a .docx file.
Run: python3 build_report.py
Output: ~/Desktop/CardioPredict_AI_Final_Report.docx
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

PROJECT = Path("/home/asifnawaz/Desktop/cardiopredict-ai")
IMG = PROJECT / "report_images"
OUT = Path.home() / "Desktop" / "CardioPredict_AI_Final_Report.docx"
doc = Document()

# ── Document-wide font: Times New Roman (academic standard) ─────── #
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)
# Apply to East Asian / complex script too so it sticks everywhere
rPr = style.element.get_or_add_rPr()
rFonts = rPr.find(qn("w:rFonts"))
if rFonts is None:
    rFonts = OxmlElement("w:rFonts")
    rPr.append(rFonts)
rFonts.set(qn("w:ascii"), "Times New Roman")
rFonts.set(qn("w:hAnsi"), "Times New Roman")
rFonts.set(qn("w:cs"), "Times New Roman")
rFonts.set(qn("w:eastAsia"), "Times New Roman")

# Force all heading styles to Times New Roman as well
for level in range(1, 5):
    try:
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Times New Roman"
        rPr = h.element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        rFonts.set(qn("w:ascii"), "Times New Roman")
        rFonts.set(qn("w:hAnsi"), "Times New Roman")
        rFonts.set(qn("w:cs"), "Times New Roman")
    except KeyError:
        pass

for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


# ──────────────────────────── helpers ──────────────────────────── #
def page_break():
    doc.add_page_break()


def H(text, level=1, align=None, color=None):
    h = doc.add_heading(text, level=level)
    if align == "center":
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if color:
        for run in h.runs:
            run.font.color.rgb = color
    return h


def P(text, bold=False, italic=False, size=None, align=None, indent=None):
    p = doc.add_paragraph()
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    return p


def bullet(text):
    doc.add_paragraph(text, style="List Bullet")


def numbered(text):
    doc.add_paragraph(text, style="List Number")


def insert_image(filename, width_cm=14, caption=None):
    """Drop an image from report_images/ centred on the page."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    img_path = IMG / filename
    if img_path.exists():
        run.add_picture(str(img_path), width=Cm(width_cm))
    else:
        run.add_text(f"[Image not found: {filename}]")
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap.add_run(caption)
        cap_run.italic = True
        cap_run.font.size = Pt(10)


def table_from_rows(headers, rows, header_bg="2563EB", widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        # background
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), header_bg)
        tcPr.append(shd)
    # Data rows
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            cell = t.rows[r].cells[c]
            cell.text = str(val)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    return t


# ════════════════════════════════════════════════════════════════ #
# COVER PAGE
# ════════════════════════════════════════════════════════════════ #
P("")
P("Final Project Report", bold=True, size=24, align="center")
P("")
P("CardioPredict AI", bold=True, size=22, align="center")
P("Heart Disease Prediction System Using Deep Learning",
  bold=True, size=16, align="center")
P("with Multi-Class Classification", bold=True, size=16, align="center")
P("")
P("")
P("Virtual University of Pakistan", bold=True, size=14, align="center")
P("")
P("")
P("")
P("Project Supervisor", bold=True, size=14, align="center")
P("Laraib Sana", size=13, align="center")
P("")
P("Submitted By", bold=True, size=14, align="center")
P("Group ID: F25PROJECT7A114", size=12, align="center")
P("")
P("Asif Nawaz Mughal     —     BC220414485",
  bold=True, size=13, align="center")
P("")
P("")
P("Software Projects & Research Section", size=12, align="center")
P("Department of Computer Sciences", size=12, align="center")
P("Virtual University of Pakistan", size=12, align="center")
P("CS619 — Fall 2025", size=11, align="center")
page_break()


# ════════════════════════════════════════════════════════════════ #
# CERTIFICATE
# ════════════════════════════════════════════════════════════════ #
P("Department of Computer Sciences", bold=True, size=16, align="center")
P("Virtual University of Pakistan", bold=True, size=16, align="center")
P("")
P("CERTIFICATE", bold=True, size=20, align="center")
P("")
P("This is to certify that Asif Nawaz Mughal (BC220414485) has worked on "
  "and completed his Software Project at Software & Research Projects "
  "Section, Department of Computer Sciences, Virtual University of "
  "Pakistan in partial fulfilment of the requirement for the degree of "
  "BS in Computer Sciences under my guidance and supervision.",
  align="justify")
P("")
P("In our opinion, it is satisfactory and up to the mark and therefore "
  "fulfils the requirements of BS in Computer Sciences.",
  align="justify")
P("")
P("")
P("Supervisor / Internal Examiner", bold=True, size=12)
P("Laraib Sana")
P("Supervisor,")
P("Software Projects & Research Section,")
P("Department of Computer Sciences,")
P("Virtual University of Pakistan")
P("")
P("_______________________")
P("(Signature)")
P("")
P("")
P("External Examiner / Subject Specialist", bold=True, size=12)
P("________________________________")
P("")
P("_______________________")
P("(Signature)")
P("")
P("Accepted By:", bold=True, align="right")
P("____________", align="right")
P("(For office use)", align="right", italic=True)
page_break()


# ════════════════════════════════════════════════════════════════ #
# EXORDIUM
# ════════════════════════════════════════════════════════════════ #
P("")
P("EXORDIUM", bold=True, size=20, align="center")
P("")
P("")
P("In the name of Allah, the Compassionate, the Merciful.",
  bold=True, align="center")
P("")
P("Praise be to Allah, Lord of Creation,", bold=True, align="center")
P("The Compassionate, the Merciful,", bold=True, align="center")
P("King of Judgment-day!", bold=True, align="center")
P("")
P("You alone we worship, and to You alone we pray for help,",
  bold=True, align="center")
P("Guide us to the straight path", bold=True, align="center")
P("")
P("The path of those who You have favored,", bold=True, align="center")
P("")
P("Not of those who have incurred Your wrath,",
  bold=True, align="center")
P("Nor of those who have gone astray.", bold=True, align="center")
page_break()


# ════════════════════════════════════════════════════════════════ #
# DEDICATION
# ════════════════════════════════════════════════════════════════ #
P("")
P("DEDICATION", bold=True, size=20, align="center")
P("")
P("")
P("")
P("Dedicated to my parents, my teachers, and my supervisor "
  "Madam Laraib Sana, whose support and guidance made this "
  "project possible.",
  italic=True, size=14, align="center")
page_break()


# ════════════════════════════════════════════════════════════════ #
# ACKNOWLEDGEMENT
# ════════════════════════════════════════════════════════════════ #
P("")
P("ACKNOWLEDGEMENT", bold=True, size=20, align="center")
P("")
P("All praise is to Allah Almighty, by Whose blessing this project "
  "reached completion.", align="justify")
P("")
P("I extend my deepest gratitude to my supervisor "
  "Madam Laraib Sana for her continuous guidance and feedback through "
  "every milestone — SRS, Design, Prototype, and Final phases.",
  align="justify")
P("")
P("I am thankful to the Software Projects & Research Section, "
  "Department of Computer Sciences, Virtual University of Pakistan "
  "for providing the platform and academic support to undertake this "
  "project as part of my BS in Computer Sciences.", align="justify")
P("")
P("Lastly, I acknowledge my family and friends for their patience, "
  "motivation, and moral support during the months of development "
  "and documentation.", align="justify")
page_break()


# ════════════════════════════════════════════════════════════════ #
# PREFACE
# ════════════════════════════════════════════════════════════════ #
P("")
P("PREFACE", bold=True, size=20, align="center")
P("")
P("This document is the Final Project Report for CardioPredict AI — "
  "Heart Disease Prediction System Using Deep Learning with Multi-Class "
  "Classification, submitted in partial fulfilment of the BS in Computer "
  "Sciences degree at Virtual University of Pakistan, under course "
  "code CS619 (Fall 2025).", align="justify")
P("")
P("Cardiovascular disease is the world's leading cause of mortality, "
  "and early identification of at-risk patients is critical. This "
  "project explores how machine learning — specifically a multi-class "
  "deep learning classifier — can support healthcare professionals by "
  "automating risk triage from routine clinical parameters.",
  align="justify")
P("")
P("The report is organised into three chapters. Chapter 1 covers "
  "requirements gathering and analysis, including functional and "
  "non-functional requirements, use cases, and the development "
  "methodology. Chapter 2 presents the system design — architecture, "
  "sequence diagrams, class diagram, database design, and user "
  "interface. Chapter 3 documents the development phase, the "
  "technology stack, the implemented modules, and screenshots from "
  "the running system.", align="justify")
page_break()


# ════════════════════════════════════════════════════════════════ #
# TABLE OF CONTENTS (placeholder — Word will regenerate)
# ════════════════════════════════════════════════════════════════ #
P("")
P("TABLE OF CONTENTS", bold=True, size=20, align="center")
P("")
P("(After opening the document in MS Word, click anywhere below this "
  "line, then go to References → Table of Contents → Automatic Table 1 "
  "to generate the TOC automatically with correct page numbers.)",
  italic=True, size=10, align="center")
P("")
P("")
P("Suggested Structure:", bold=True)
toc_lines = [
    "Chapter 1   Gathering & Analyzing Info",
    "    1.1   Introduction",
    "    1.2   Purpose",
    "    1.3   Scope",
    "    1.4   Definitions, Acronyms and Abbreviations",
    "    1.5   Project Requirements",
    "        1.5.1   Functional Requirements",
    "        1.5.2   Non-Functional Requirements",
    "    1.6   Use Cases and Usage Scenarios",
    "        1.6.1   Use Case Diagram",
    "        1.6.2   Usage Scenarios",
    "    1.7   Development Methodology",
    "        1.7.1   Chosen Methodology",
    "        1.7.2   Reasons for Chosen Methodology",
    "        1.7.3   Work Plan (Gantt Chart)",
    "        1.7.4   Project Schedule",
    "",
    "Chapter 2   Designing the Project",
    "    2.1   Introduction",
    "    2.2   Purpose",
    "    2.3   Scope",
    "    2.4   Definitions, Acronyms and Abbreviations",
    "    2.5   Architectural Representation",
    "    2.6   Dynamic Model: Sequence Diagrams",
    "    2.7   Object Model / Logical Model: Class Diagram",
    "    2.8   Database Model",
    "    2.9   Graphical User Interfaces",
    "",
    "Chapter 3   Development",
    "    3.1   Introduction",
    "    3.2   Development Plan / Architecture",
    "    3.3   Technology Stack",
    "    3.4   Implemented Modules",
    "    3.5   Machine Learning Pipeline",
    "    3.6   Performance Results",
    "    3.7   Database Implementation",
    "    3.8   Screenshots of the Running Application",
    "    3.9   Conclusion",
    "",
    "References",
    "Appendix",
]
for line in toc_lines:
    P(line, size=11)
page_break()


# ════════════════════════════════════════════════════════════════ #
# CHAPTER 1 TITLE PAGE
# ════════════════════════════════════════════════════════════════ #
P("")
P("")
P("")
P("")
P("CHAPTER 1", bold=True, size=28, align="center")
P("Gathering & Analyzing Info", bold=True, size=20, align="center")
page_break()


# ──────── 1.1 Introduction ──────── #
H("1.  Gathering & Analyzing Info", level=1)
H("1.1   Introduction", level=2)
P("This chapter presents the requirements analysis phase for "
  "CardioPredict AI, a web-based clinical decision support system "
  "that classifies heart disease risk into three categories — "
  "Low, Medium, and High — using a deep learning model. The "
  "chapter defines the project's purpose and scope, the "
  "functional and non-functional requirements gathered during "
  "the SRS phase, the principal use cases and usage scenarios, "
  "and the development methodology adopted to deliver the system "
  "within the academic timeline.", align="justify")

# 1.2 Purpose
H("1.2   Purpose", level=2)
P("The purpose of CardioPredict AI is to assist healthcare "
  "professionals in early identification of patients at risk of "
  "heart disease. The system accepts standard clinical "
  "parameters such as age, gender, blood pressure, cholesterol, "
  "ECG results, and other features, processes them through "
  "trained machine learning models (Artificial Neural Network, "
  "Logistic Regression, and Random Forest), and returns a "
  "multi-class risk classification with associated probability "
  "scores. By providing instant, evidence-based risk assessment "
  "alongside downloadable clinical reports, the system aims to "
  "reduce diagnostic delays and support timely intervention "
  "decisions.", align="justify")

# 1.3 Scope
H("1.3   Scope", level=2)
P("The main objective of this project is to design and develop "
  "a heart disease prediction system using deep learning and "
  "machine learning techniques. The system analyses patient "
  "medical data — including age, gender, blood pressure, "
  "cholesterol levels, ECG results, and other clinical "
  "indicators — to predict heart disease risk and classify "
  "patients into three risk categories: Low Risk, Medium Risk, "
  "and High Risk.", align="justify")
P("")
P("Key Components:", bold=True)
bullet("Machine Learning Models: Artificial Neural Networks (ANN), "
       "Logistic Regression, and Random Forest Classifier — "
       "trained and compared on the UCI Heart Disease dataset.")
bullet("Advanced Techniques: Dropout regularisation, L2 "
       "regularisation, and hyperparameter tuning to optimise "
       "model performance and prevent overfitting.")
bullet("Web-Based Interface: A user-friendly web application for "
       "healthcare professionals to input patient data and "
       "receive instant risk predictions with probability scores.")
bullet("Data Management: Secure database storage for patient "
       "history, prediction results, and visualisation of risk "
       "trends over time.")
bullet("Report Generation: Ability to download detailed "
       "prediction reports in PDF format and history records "
       "in CSV format.")
P("")
P("The final deliverable is a working clinical decision support "
  "system that supports early detection of heart disease, "
  "reduces human error in diagnosis, and provides clinical "
  "decision support to healthcare professionals.",
  align="justify")

# 1.4 Definitions, Acronyms
H("1.4   Definitions, Acronyms and Abbreviations", level=2)
table_from_rows(
    headers=["Term", "Definition"],
    rows=[
        ["ANN", "Artificial Neural Network"],
        ["API", "Application Programming Interface"],
        ["CRUD", "Create, Read, Update, Delete"],
        ["CS619", "Final Year Project course code (Virtual University of Pakistan)"],
        ["CSV", "Comma-Separated Values"],
        ["EHR", "Electronic Health Record"],
        ["ERD", "Entity Relationship Diagram"],
        ["FR / NFR", "Functional / Non-Functional Requirement"],
        ["JWT", "JSON Web Token"],
        ["LR", "Logistic Regression"],
        ["ML / DL", "Machine Learning / Deep Learning"],
        ["ORM", "Object-Relational Mapping"],
        ["PDF", "Portable Document Format"],
        ["RF", "Random Forest"],
        ["ROC-AUC", "Receiver Operating Characteristic — Area Under Curve"],
        ["SRS", "Software Requirements Specification"],
        ["SQL", "Structured Query Language"],
        ["UCI", "University of California, Irvine (dataset source)"],
        ["UI / GUI", "User Interface / Graphical User Interface"],
        ["UML", "Unified Modeling Language"],
        ["VU", "Virtual University of Pakistan"],
    ],
    widths=[3.5, 12],
)

# 1.5 Project Requirements
H("1.5   Project Requirements", level=2)

H("1.5.1   Functional Requirements", level=3)
fr_data = [
    ("FR-01: Patient Data Input",
     "The system shall allow healthcare professionals to enter "
     "patient medical information including age, gender, blood "
     "pressure, cholesterol level, blood sugar, ECG results, "
     "and other relevant clinical parameters through an "
     "interactive web form."),
    ("FR-02: Data Preprocessing",
     "The system shall automatically preprocess entered data by "
     "handling missing values, normalising numerical fields, "
     "and encoding categorical variables before sending to "
     "prediction models."),
    ("FR-03: Heart Disease Prediction",
     "The system shall analyse processed patient data using "
     "trained machine learning and deep learning models to "
     "predict the risk level of heart disease."),
    ("FR-04: Multi-Class Risk Classification",
     "The system shall classify patients into three categories: "
     "Low Risk, Medium Risk, and High Risk based on model "
     "predictions and probability scores."),
    ("FR-05: Probability Display",
     "The system shall display probability percentages for each "
     "risk category to help healthcare professionals understand "
     "the confidence level of predictions."),
    ("FR-06: Model Comparison",
     "The system shall utilise multiple algorithms (ANN, "
     "Logistic Regression, Random Forest) and compare their "
     "performance metrics for accuracy and reliability."),
    ("FR-07: Visualisation of Results",
     "The system shall generate visual charts and graphs "
     "representing risk trends and prediction results for "
     "better understanding and clinical communication."),
    ("FR-08: Patient History Storage",
     "The system shall store patient data and past prediction "
     "results in a secure database for record-keeping and "
     "future reference."),
    ("FR-09: Report Generation",
     "The system shall allow users to download detailed "
     "prediction reports in PDF format containing patient data, "
     "prediction results, and probability scores."),
    ("FR-10: User Interface (Web-Based GUI)",
     "The system shall provide a simple, intuitive, and "
     "user-friendly interface that allows healthcare "
     "professionals to easily interact with the system without "
     "technical training."),
]
for fr_id, fr_desc in fr_data:
    p = doc.add_paragraph()
    r = p.add_run(fr_id + "  ")
    r.bold = True
    p.add_run(fr_desc)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

H("1.5.2   Non-Functional Requirements", level=3)
nfr_data = [
    ("NFR-01: Performance",
     "The system shall generate prediction results within 3–5 "
     "seconds after data submission to ensure smooth user "
     "experience and clinical usability."),
    ("NFR-02: Accuracy",
     "Target accuracy ≥ 75% (achieved 77.5% on the UCI Cleveland "
     "+ Hungarian test set with the ANN model)."),
    ("NFR-03: Usability",
     "The interface shall be intuitive and easy to understand, "
     "with real-time form validation, inline status indicators, "
     "and toast notifications. No technical training required."),
    ("NFR-04: Reliability",
     "The system shall perform consistently without crashes or "
     "incorrect outputs during normal operation, designed for "
     "99% uptime in production environments."),
    ("NFR-05: Security",
     "Passwords stored using bcrypt hashing; sessions secured "
     "with JWT-signed tokens; SQL injection prevented via "
     "SQLAlchemy parameterised queries; HTTPS deployment "
     "recommended for production."),
    ("NFR-06: Scalability",
     "Modular architecture (Next.js + FastAPI + MySQL) designed "
     "to scale; for FYP scope, system is validated for "
     "single-doctor demonstration."),
    ("NFR-07: Maintainability",
     "Modular codebase with reusable components, separated "
     "controllers / routes / models, and SQLAlchemy ORM "
     "facilitating schema evolution."),
    ("NFR-08: Compatibility",
     "Tested on Chrome, Firefox, and Edge (latest versions); "
     "runs on Windows, macOS, and Linux."),
    ("NFR-09: Availability",
     "Production architecture supports 24/7 operation with "
     "database backups; FYP deployment runs on local "
     "environment."),
    ("NFR-10: Data Integrity",
     "Foreign-key constraints, transactional commits, and "
     "Pydantic schema validation guarantee data consistency."),
]
for nfr_id, nfr_desc in nfr_data:
    p = doc.add_paragraph()
    r = p.add_run(nfr_id + "  ")
    r.bold = True
    p.add_run(nfr_desc)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# 1.6 Use Cases
H("1.6   Use Cases and Usage Scenarios", level=2)

H("1.6.1   Use Case Diagram", level=3)
P("The following actors interact with the system:", align="justify")
bullet("Healthcare Professional (Doctor) — Primary user who "
       "inputs patient data, views predictions, generates "
       "reports, and manages patient history.")
bullet("System Administrator — Manages system configuration, "
       "user accounts, and patient reports.")
P("")
insert_image("image6.png", width_cm=14,
             caption="Figure 1.1 — Use Case Diagram of CardioPredict AI")
P("")
P("Major use cases supported by the system:", align="justify")
for uc in ["Manual Patient Entry", "Input Health Parameters",
           "Run Heart Disease Prediction",
           "View Prediction & Probabilities",
           "Generate PDF Report", "View Prediction History",
           "Export History as CSV", "Manage Users (Admin)",
           "View All Patients (Admin)",
           "View Patient Full Report (Admin)"]:
    bullet(uc)
P("")
insert_image("image8.png", width_cm=14,
             caption="Figure 1.3 — Context Diagram showing system scope and external entities")

H("1.6.2   Usage Scenarios", level=3)

# Use case tables
USE_CASES = [
    ("UC-01: Manual Patient Entry", [
        ("Use Case ID", "UC-01"),
        ("Actor", "Healthcare Professional"),
        ("Description", "User manually enters patient demographics and clinical features for prediction."),
        ("Pre-Condition", "User must be logged in with valid credentials."),
        ("Post-Condition", "Data passes to preprocessing module for analysis."),
        ("Main Flow",
         "1. User navigates to patient entry form.\n"
         "2. Enters patient demographics (name, age, gender, contact).\n"
         "3. Enters clinical parameters (BP, cholesterol, etc.).\n"
         "4. System validates input client-side.\n"
         "5. User submits form.\n"
         "6. Data forwarded to preprocessing."),
        ("Alternative Paths",
         "Alt 1: Missing required fields → System shows validation errors.\n"
         "Alt 2: Invalid data format → System displays correction guidance."),
        ("Exceptions",
         "Exc 1: Network failure → Error displayed, retry option offered."),
    ]),
    ("UC-02: Run Heart Disease Prediction", [
        ("Use Case ID", "UC-11"),
        ("Actor", "Healthcare Professional"),
        ("Description", "Doctor submits clinical features; system returns Low/Medium/High risk classification with probabilities."),
        ("Pre-Condition", "Trained model loaded; patient data filled in."),
        ("Post-Condition", "Prediction stored in database with risk class."),
        ("Main Flow",
         "1. User clicks 'Run Prediction'.\n"
         "2. Frontend sends POST /predictions/predict with health data.\n"
         "3. Backend validates schema (Pydantic).\n"
         "4. ML engine scales features and runs ANN inference.\n"
         "5. System computes Low/Medium/High probabilities.\n"
         "6. Prediction record saved to database.\n"
         "7. Result page displays risk class with confidence."),
        ("Alternative Paths",
         "Alt 1: User selects different model (LR, RF) → System runs inference with selected model."),
        ("Exceptions",
         "Exc 1: Model file not loaded → Backend returns 503 with 'run training first' message."),
    ]),
    ("UC-03: Generate Downloadable Report", [
        ("Use Case ID", "UC-18"),
        ("Actor", "Healthcare Professional"),
        ("Description", "User generates a PDF report of a prediction with patient info, risk class, probabilities, and recommendations."),
        ("Pre-Condition", "Prediction completed and result loaded on /results/[id] page."),
        ("Post-Condition", "PDF file saved to user's Downloads folder."),
        ("Main Flow",
         "1. User clicks 'Export PDF' button.\n"
         "2. Frontend dynamically loads jsPDF library.\n"
         "3. PDF document is built with header, risk card, "
         "probability table, parameters, recommendations, footer.\n"
         "4. doc.save() triggers browser download.\n"
         "5. PDF saved with filename CardioPredict_Report_<id>_<patient>.pdf."),
        ("Alternative Paths",
         "Alt 1: User clicks Print instead → browser print dialog opens."),
        ("Exceptions",
         "Exc 1: jsPDF load failure → Error toast displayed."),
    ]),
    ("UC-04: View Prediction History", [
        ("Use Case ID", "UC-19"),
        ("Actor", "Healthcare Professional"),
        ("Description", "User views all past predictions with filters and exports them as CSV."),
        ("Pre-Condition", "User logged in; predictions exist in database."),
        ("Post-Condition", "User views or downloads prediction history."),
        ("Main Flow",
         "1. User navigates to /history page.\n"
         "2. Frontend fetches predictions from /predictions/history.\n"
         "3. Table displays risk badges, confidence, model, date.\n"
         "4. User filters by patient name or risk class.\n"
         "5. User clicks 'Export CSV' → file downloaded."),
        ("Alternative Paths",
         "Alt 1: No predictions yet → Empty-state message shown."),
        ("Exceptions",
         "Exc 1: Network error → Toast displayed, retry available."),
    ]),
    ("UC-05: Manage Users (Admin)", [
        ("Use Case ID", "UC-A1"),
        ("Actor", "System Administrator"),
        ("Description", "Admin manages user accounts: promotes to admin, disables, or deletes users."),
        ("Pre-Condition", "Admin user logged into /admin panel."),
        ("Post-Condition", "User record updated in database."),
        ("Main Flow",
         "1. Admin opens /admin/user/list.\n"
         "2. Selects a user row.\n"
         "3. Clicks Edit / Delete / Promote.\n"
         "4. SQLAdmin updates record.\n"
         "5. Database transaction commits."),
        ("Alternative Paths",
         "Alt 1: Admin disables their own account → System warns but allows."),
        ("Exceptions",
         "Exc 1: Foreign-key constraint violation → SQLAdmin shows error."),
    ]),
]

for uc_title, uc_rows in USE_CASES:
    H(uc_title, level=4)
    table_from_rows(
        headers=["Attribute", "Details"],
        rows=uc_rows,
        widths=[4, 11.5],
        header_bg="475569",
    )
    P("")

P("Note: Additional use cases (UC-02 Import Dataset, UC-03 to "
  "UC-05 Preprocessing, UC-06 to UC-08 Model Training, UC-09 "
  "Regularisation, UC-10 Cross-Validation) were also documented "
  "in the SRS for the offline / training-time phase. Their full "
  "details are available in the Software Requirements "
  "Specification document submitted on 04 Dec 2025.",
  italic=True, align="justify")

# 1.7 Methodology
H("1.7   Development Methodology", level=2)

H("1.7.1   Chosen Methodology", level=3)
P("The methodology adopted is the Iterative & Incremental "
  "Development Model. Each major module of the system — "
  "authentication, patient management, ML prediction, results, "
  "history, and admin panel — was developed and tested in a "
  "separate iteration, with each iteration delivering working "
  "software integrated with the previous output.",
  align="justify")
P("")
insert_image("image5.png", width_cm=12,
             caption="Figure 1.4 — Iterative Development Model adopted for CardioPredict AI")

H("1.7.2   Reasons for Chosen Methodology", level=3)
bullet("The system has clearly separable functional modules, "
       "perfect for incremental delivery.")
bullet("Working software at the end of every iteration enables "
       "continuous testing.")
bullet("The ML pipeline can be refined as more data and features "
       "are added without restarting.")
bullet("Bugs are caught early instead of at the end of the project.")
bullet("Flexibility to add new features without rewrites.")

H("1.7.3   Work Plan (Gantt Chart)", level=3)
P("The project ran from November 2025 to May 2026 across four "
  "official VU phases:", align="justify")
table_from_rows(
    headers=["Phase", "Start", "End", "Duration", "Deliverable"],
    rows=[
        ["1 — SRS", "05 Nov 2025", "04 Dec 2025", "30 days", "SRS Document"],
        ["2 — Design", "05 Dec 2025", "23 Jan 2026", "50 days", "Design Document"],
        ["3 — Prototype", "24 Jan 2026", "03 Mar 2026", "40 days", "Working prototype"],
        ["4 — Final Deliverable", "04 Mar 2026", "11 May 2026", "70 days", "Final system + report"],
    ],
    widths=[3, 2.5, 2.5, 2, 4],
)
P("")
insert_image("image10.png", width_cm=15,
             caption="Figure 1.2 — Project Schedule (Gantt Chart) Nov 2025 → May 2026")

H("1.7.4   Project Schedule (Submission Calendar)", level=3)
P("The project followed VU's official submission calendar: SRS "
  "submitted on 04 Dec 2025, Design Document on 23 Jan 2026, "
  "Prototype phase ending 03 Mar 2026, and Final Deliverable "
  "submitted on 11 May 2026. All four milestones were met on or "
  "before the deadline.", align="justify")
page_break()


# ════════════════════════════════════════════════════════════════ #
# CHAPTER 2 TITLE PAGE
# ════════════════════════════════════════════════════════════════ #
P("")
P("")
P("")
P("")
P("CHAPTER 2", bold=True, size=28, align="center")
P("Designing the Project", bold=True, size=20, align="center")
page_break()

H("2.  Designing the Project", level=1)
H("2.1   Introduction", level=2)
P("This chapter presents the system design of CardioPredict AI. "
  "It documents the high-level three-tier architecture that "
  "separates presentation, application, and data concerns; the "
  "dynamic behaviour of the system through sequence diagrams; "
  "the static object-oriented structure through the class "
  "diagram; the physical database model in MySQL; and "
  "screenshots of the principal user interfaces.", align="justify")

H("2.2   Purpose", level=2)
P("The purpose of this chapter is to translate the requirements "
  "documented in Chapter 1 into a concrete technical design that "
  "guided the implementation phase. The design serves as a "
  "contract between the requirements and the eventual code.",
  align="justify")

H("2.3   Scope", level=2)
P("The design covers all subsystems of CardioPredict AI: the "
  "Next.js frontend, the FastAPI backend, the MySQL database, "
  "the ML prediction service, and the admin panel. External "
  "actors (Doctor, Administrator) and secondary actors (ML "
  "Engine, Database) are included.", align="justify")

H("2.4   Definitions, Acronyms and Abbreviations", level=2)
P("Refer to §1.4 for the master glossary. Additional terms used "
  "in this chapter:", align="justify")
bullet("MVC — Model-View-Controller pattern.")
bullet("REST — Representational State Transfer.")
bullet("DAO — Data Access Object.")
bullet("FK — Foreign Key.")
bullet("PK — Primary Key.")

H("2.5   Architectural Representation", level=2)
P("CardioPredict AI follows a Three-Tier Architecture ensuring "
  "separation of concerns, scalability, and maintainability:",
  align="justify")
table_from_rows(
    headers=["Tier", "Components", "Technology"],
    rows=[
        ["Presentation", "Web UI, Forms, Charts", "Next.js 14, Tailwind CSS, Chart.js"],
        ["Application", "REST API, Auth, ML Service", "FastAPI, JWT, TensorFlow"],
        ["Data", "Database, ORM, File Storage", "MySQL 8, SQLAlchemy, Local FS"],
    ],
    widths=[3, 5, 7],
)
P("")
insert_image("image7.png", width_cm=15,
             caption="Figure 2.1 — Three-Tier Architecture of CardioPredict AI")

H("2.6   Dynamic Model: Sequence Diagrams", level=2)
P("Three principal sequence diagrams describe the runtime "
  "behaviour of the system:", align="justify")
P("")
P("Phase 1 — Manual Patient Entry", bold=True)
P("Doctor enters patient demographics and clinical parameters; "
  "the backend validates input, stores the patient and "
  "health-parameter records, and confirms the new Patient ID.",
  align="justify")
insert_image("image11.png", width_cm=15,
             caption="Figure 2.2 — Sequence Diagram: Manual Patient Entry")
P("")
P("Phase 4 — Online Prediction", bold=True)
P("Doctor selects a patient and submits clinical data; the "
  "Prediction API loads the trained ML model, computes "
  "Low/Medium/High risk probabilities, persists the result, "
  "and returns it to the dashboard.", align="justify")
insert_image("image12.png", width_cm=15,
             caption="Figure 2.3 — Sequence Diagram: Online Prediction (Run Risk Assessment)")
P("")
P("Phase 5 — Download and Export", bold=True)
P("Doctor clicks 'Export PDF'; the frontend dynamically loads "
  "jsPDF, builds the document client-side, and triggers a "
  "browser download.", align="justify")
insert_image("image13.png", width_cm=15,
             caption="Figure 2.4 — Sequence Diagram: Report Download & Export")

H("2.7   Object Model / Logical Model: Class Diagram", level=2)
P("The system is designed using Object-Oriented Programming "
  "principles. The following classes form the domain model:",
  align="justify")
table_from_rows(
    headers=["Class", "Purpose"],
    rows=[
        ["User", "Stores doctors and administrators with credentials and role."],
        ["Patient", "Patient demographics linked to a doctor."],
        ["HealthParameter", "Stores 13 clinical features per record."],
        ["Prediction", "Risk classification result with probabilities and FK references."],
        ["MLModel", "Registry of trained ML models with performance metrics."],
    ],
    widths=[4, 11],
)
P("")
insert_image("image14.png", width_cm=15,
             caption="Figure 2.5 — Class Diagram (Object / Logical Model)")

H("2.8   Database Model", level=2)
P("CardioPredict AI uses MySQL 8 with InnoDB. Five tables "
  "connected through six 1:N relationships make up the schema:",
  align="justify")
table_from_rows(
    headers=["Table", "Primary Key", "Key Foreign Keys"],
    rows=[
        ["users", "id", "—"],
        ["patients", "id", "doctor_id → users(id)"],
        ["health_parameters", "id", "patient_id → patients(id)"],
        ["predictions", "id",
         "patient_id, health_parameter_id, model_id, user_id"],
        ["ml_models", "id", "—"],
    ],
    widths=[4, 3, 8.5],
)
P("")
insert_image("image15.png", width_cm=15,
             caption="Figure 2.6 — Entity-Relationship Diagram")
P("")
insert_image("image16.png", width_cm=15,
             caption="Figure 2.7 — Database Design (Physical Schema)")

H("2.9   Graphical User Interfaces", level=2)
P("The following major interfaces are designed:", align="justify")
bullet("Login & Registration page — JWT authentication with "
       "real-time username / email availability checks.")
bullet("Doctor Dashboard — Stats cards plus recent predictions list.")
bullet("New Prediction Form — 13 clinical inputs with template "
       "buttons for sample data.")
bullet("Result Page — Colour-coded risk banner, charts, "
       "recommendations, Export PDF button.")
bullet("History Page — Searchable / filterable table of past "
       "predictions with CSV export.")
bullet("Admin Panel — User and patient management with role-based "
       "access, mounted at /admin.")
P("")
P("Screenshots are presented in Chapter 3 § 3.8.",
  italic=True, align="justify")
page_break()


# ════════════════════════════════════════════════════════════════ #
# CHAPTER 3 TITLE PAGE
# ════════════════════════════════════════════════════════════════ #
P("")
P("")
P("")
P("")
P("CHAPTER 3", bold=True, size=28, align="center")
P("Development", bold=True, size=20, align="center")
page_break()

H("3.  Development", level=1)

H("3.1   Introduction", level=2)
P("This chapter documents the implementation of CardioPredict "
  "AI, including the technology stack chosen, the modules "
  "built, the machine learning pipeline, and screenshots of "
  "the running application. Where the implementation differs "
  "from the original Design Document, those decisions are "
  "explained.", align="justify")

H("3.2   Development Plan / Architecture", level=2)
P("CardioPredict AI follows a three-tier architecture:",
  align="justify")
bullet("Presentation tier — Next.js 14 (App Router) with React "
       "18, Tailwind CSS for styling, Chart.js for "
       "risk-distribution charts, and react-hot-toast for user "
       "feedback.")
bullet("Application tier — FastAPI (Python 3.10) exposing REST "
       "endpoints under /auth, /patients, /predictions. JWT "
       "authentication, Pydantic request validation, and "
       "SQLAlchemy ORM.")
bullet("Data tier — MySQL 8 (InnoDB) accessed via PyMySQL "
       "driver. Trained ML models are persisted as .keras and "
       ".pkl files in the file system.")
P("")
P("The frontend communicates with the backend over JSON-formatted "
  "HTTP requests carrying a JWT bearer token. CORS is configured "
  "on the backend to allow the Next.js dev server "
  "(http://localhost:3000) to call the API "
  "(http://localhost:8000).", align="justify")

H("3.3   Technology Stack", level=2)
table_from_rows(
    headers=["Category", "Tools / Technologies Used"],
    rows=[
        ["IDE & Editors", "Visual Studio Code, MySQL Workbench"],
        ["Frontend Framework", "Next.js 14, React 18, Tailwind CSS"],
        ["Frontend Libraries",
         "Chart.js, Lucide React, jsPDF, jspdf-autotable, react-hot-toast"],
        ["Backend Framework", "Python 3.10, FastAPI, Uvicorn"],
        ["Backend Libraries",
         "SQLAlchemy 2.0, Pydantic, python-jose, passlib, bcrypt, SQLAdmin"],
        ["Database", "MySQL 8 (InnoDB), PyMySQL"],
        ["Machine Learning",
         "TensorFlow / Keras, scikit-learn, pandas, numpy, matplotlib, joblib"],
        ["Version Control", "Git, GitHub"],
        ["API Testing", "Postman, FastAPI Swagger UI (/docs)"],
        ["Diagrams", "draw.io / diagrams.net, Lucidchart"],
        ["Documentation", "Microsoft Word, OnlyOffice Impress"],
        ["Operating System", "Ubuntu Linux 24.04"],
    ],
    widths=[5, 10.5],
)

H("3.4   Implemented Modules", level=2)

H("3.4.1   Authentication Module", level=3)
P("JWT-based login using bcrypt for password hashing. The "
  "registration form performs real-time username and email "
  "availability checks via debounced calls to "
  "/auth/check-availability. Live password-strength rules "
  "enforce a minimum of 8 characters with at least one "
  "uppercase, one lowercase, and one digit. Login accepts "
  "either username or email.", align="justify")

H("3.4.2   Patient Management Module", level=3)
P("CRUD operations on the patients table. Doctors can search, "
  "add, edit, and delete patients. New patients can also be "
  "created on-the-fly during prediction submission if no "
  "existing patient is selected, simplifying the workflow.",
  align="justify")

H("3.4.3   ML Prediction Module", level=3)
P("The /predictions/predict endpoint receives a feature vector, "
  "scales it with the saved StandardScaler, runs inference "
  "through the selected model (ANN by default), persists the "
  "Prediction row, and returns the risk class together with "
  "the three probability scores.", align="justify")

H("3.4.4   Results & Reporting Module", level=3)
P("A dedicated results page renders a colour-coded risk banner, "
  "a doughnut chart of probabilities, a comparison bar chart, "
  "clinical recommendations, and the input parameters used. "
  "Reports can be exported to PDF entirely on the client using "
  "jsPDF + jspdf-autotable — no server round-trip is required, "
  "making exports near-instantaneous.", align="justify")

H("3.4.5   History Module", level=3)
P("The history page lists all predictions across patients with "
  "risk-class filters and a search box. The list can be "
  "exported as a CSV with heart.csv-compatible columns "
  "(age, sex, cp, trestbps, …, thal, risk_class, confidence, "
  "model_used, predicted_at), enabling data scientists to "
  "re-feed the records into training pipelines if needed.",
  align="justify")

H("3.4.6   Admin Panel", level=3)
P("Built with SQLAdmin mounted at /admin, gated by role-based "
  "authentication. Admins can manage Users (promote, disable), "
  "view all Patients across doctors, and download per-patient "
  "HTML reports via a custom /reports/patient/{id} route. The "
  "default admin account (admin / Admin@321) is bootstrapped by "
  "the create_admin.py script.", align="justify")

H("3.5   Machine Learning Pipeline", level=2)
P("Dataset:", bold=True)
P("UCI Heart Disease — Cleveland (303 records) + Hungarian (294 "
  "records), combined to ~597 rows after preprocessing.",
  align="justify")
P("")
P("Preprocessing:", bold=True)
bullet("Missing values (?) replaced with NaN and imputed with column median.")
bullet("Multi-class target derived from original 0–4 score: 0 → Low, 1–2 → Medium, 3–4 → High.")
bullet("Features standardised with StandardScaler (mean = 0, std = 1).")
bullet("80 / 20 train-test split with stratified sampling.")
P("")
P("Models trained:", bold=True)
bullet("ANN (TensorFlow / Keras) — 3 hidden layers (128 → 64 → "
       "32) with ReLU activations, BatchNormalisation, Dropout "
       "(0.3), L2 regularisation (0.001), softmax output, Adam "
       "optimiser, EarlyStopping with patience 15.")
bullet("Logistic Regression (scikit-learn) — lbfgs solver, "
       "multi-class, C = 1.0.")
bullet("Random Forest (scikit-learn) — 100 estimators, max depth 10.")

H("3.6   Performance Results", level=2)
table_from_rows(
    headers=["Model", "Accuracy", "F1 Score", "ROC-AUC"],
    rows=[
        ["ANN (Deep Learning)", "77.5%", "0.77", "0.917"],
        ["Logistic Regression", "76.7%", "0.76", "0.910"],
        ["Random Forest", "76.7%", "0.75", "0.905"],
    ],
    widths=[5, 3, 3, 3],
)
P("")
P("The ANN was selected as the primary model on the basis of "
  "the highest ROC-AUC (0.917), indicating the best "
  "class-discrimination ability across Low / Medium / High "
  "categories.", align="justify")

H("3.7   Database Implementation", level=2)
P("Five tables — users, patients, health_parameters, "
  "predictions, ml_models — are created automatically on "
  "application start through Base.metadata.create_all() from "
  "SQLAlchemy. Foreign keys enforce referential integrity. "
  "Auto-incrementing primary keys, ENUM constraints (e.g. "
  "risk_class), and DEFAULT NOW() timestamps are configured at "
  "the database level rather than purely in application code.",
  align="justify")

H("3.8   Screenshots of the Running Application", level=2)
P("The following screenshots show the principal interfaces of "
  "CardioPredict AI captured from the running prototype on "
  "localhost:3000.", align="justify")
P("")

screenshots = [
    ("image17.png", "Figure 3.1 — Login Page (with real-time validation)"),
    ("image18.png", "Figure 3.2 — Registration Page (Step 1: Account; Step 2: Doctor Profile)"),
    ("image19.png", "Figure 3.3 — Doctor Dashboard (stats cards + recent predictions)"),
    ("image20.png", "Figure 3.4 — New Prediction Form (13 clinical inputs + template buttons)"),
    ("image21.png", "Figure 3.5 — Prediction Result Page (risk banner, charts, recommendations)"),
]
for fname, caption in screenshots:
    insert_image(fname, width_cm=15, caption=caption)
    P("")

H("3.9   Conclusion", level=2)
P("The CardioPredict AI system was successfully developed across "
  "five iterations, delivering authentication, patient "
  "management, three-model ML prediction, reporting, and an "
  "admin panel. The best ANN model achieved 77.5% accuracy and "
  "0.917 ROC-AUC on the UCI dataset. All twenty test cases "
  "documented in the Design Document passed with no critical "
  "defects.", align="justify")
P("")
P("Future Work:", bold=True)
bullet("Train on larger multi-hospital datasets for "
       "clinical-grade accuracy (>90%).")
bullet("Integrate with hospital EHR systems via FHIR APIs.")
bullet("Add explainable AI techniques (SHAP / LIME) to "
       "highlight which features drove each prediction.")
bullet("Deploy to a cloud platform with HTTPS and email "
       "notifications for risk alerts.")
P("")
P("Sincere thanks to my supervisor Madam Laraib Sana for her "
  "continuous guidance throughout this project.",
  italic=True, align="justify")
page_break()


# ════════════════════════════════════════════════════════════════ #
# REFERENCES
# ════════════════════════════════════════════════════════════════ #
H("REFERENCES", level=1, align="center")
references = [
    "R. Detrano et al., \"International application of a new "
    "probability algorithm for the diagnosis of coronary artery "
    "disease,\" American Journal of Cardiology, vol. 64, "
    "pp. 304–310, 1989.",

    "UCI Machine Learning Repository — Heart Disease Data Set, "
    "https://archive.ics.uci.edu/ml/datasets/heart+disease "
    "(accessed Apr 2026).",

    "World Health Organization, \"Cardiovascular Diseases (CVDs) "
    "— Fact Sheet,\" https://www.who.int/news-room/fact-sheets/"
    "detail/cardiovascular-diseases-(cvds)",

    "F. Chollet, Deep Learning with Python, 2nd ed. "
    "Manning Publications, 2021.",

    "A. Géron, Hands-On Machine Learning with Scikit-Learn, "
    "Keras, and TensorFlow, 3rd ed. O'Reilly Media, 2022.",

    "FastAPI Documentation, https://fastapi.tiangolo.com "
    "(accessed Apr 2026).",

    "Next.js 14 Documentation, https://nextjs.org/docs "
    "(accessed Apr 2026).",

    "SQLAlchemy 2.0 Documentation, https://docs.sqlalchemy.org "
    "(accessed Apr 2026).",

    "TensorFlow Keras API, https://www.tensorflow.org/api_docs/"
    "python/tf/keras (accessed Apr 2026).",

    "scikit-learn User Guide, https://scikit-learn.org/stable/"
    "user_guide.html (accessed Apr 2026).",
]
for i, ref in enumerate(references, start=1):
    p = doc.add_paragraph(style="List Number")
    p.add_run(ref)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
page_break()


# ════════════════════════════════════════════════════════════════ #
# APPENDIX
# ════════════════════════════════════════════════════════════════ #
H("APPENDIX", level=1, align="center")

H("Appendix A — Default Admin Credentials", level=2)
P("After running python create_admin.py from the backend "
  "directory, the system bootstraps a default administrator "
  "account with the following credentials:", align="justify")
table_from_rows(
    headers=["Field", "Value"],
    rows=[
        ["URL", "http://localhost:8000/admin"],
        ["Username", "admin"],
        ["Email", "admin@cardioai.com"],
        ["Password", "Admin@321"],
    ],
    widths=[4, 9],
)
P("")
P("This password should be changed before any production "
  "deployment by editing the ADMIN_PASSWORD constant in "
  "backend/create_admin.py and re-running the script.",
  italic=True, align="justify")
P("")

H("Appendix B — Source Code Repository", level=2)
P("The complete source code for CardioPredict AI is available at:",
  align="justify")
P("https://github.com/asifnawazmughal/cardiopredict-ai", bold=True)
P("")
P("Repository structure:", bold=True)
P("backend/             — FastAPI server, ML pipeline, SQLAdmin", indent=0.5)
P("heart-disease-prediction/   — Next.js frontend", indent=0.5)
P("README.md            — Setup and run instructions", indent=0.5)
P("")

H("Appendix C — Performance Test Run (Sample)", level=2)
P("Trained models output on Cleveland + Hungarian dataset:",
  align="justify")
P("")
P("Model Comparison", bold=True)
P("─" * 40, italic=True)
P("ANN:                 Accuracy 0.7750  F1 0.7745  ROC-AUC 0.9168", italic=True)
P("LogisticRegression:  Accuracy 0.7667  F1 0.7638  ROC-AUC 0.9095", italic=True)
P("RandomForest:        Accuracy 0.7667  F1 0.7462  ROC-AUC 0.9046", italic=True)
P("─" * 40, italic=True)
P("All models saved in: backend/ml/saved_models/", italic=True)


# ════════════════════════════════════════════════════════════════ #
# SAVE
# ════════════════════════════════════════════════════════════════ #
doc.save(OUT)
print(f"✅ Report saved to: {OUT}")
print(f"   File size: {OUT.stat().st_size / 1024:.1f} KB")
